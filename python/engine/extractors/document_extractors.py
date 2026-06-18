"""Document extractors for PDF and DOCX parsing."""

import io
from typing import Optional

from ..base import BaseExtractor, ScrapeResult


class PdfExtractor(BaseExtractor):
    name = "pdf_extractor"

    def extract(self, raw_bytes: bytes, result: ScrapeResult, content_type: str = "") -> ScrapeResult:
        """Extract text from PDF bytes, appending markdown to the result."""
        if not raw_bytes or content_type not in ("application/pdf", "pdf", ""):
            return result

        text_parts: list[str] = []

        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except ImportError:
            pass

        if not text_parts:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            except ImportError:
                pass

        if text_parts:
            markdown = "\n\n".join(text_parts)
            result.markdown = markdown
            result.text = markdown
            result.metadata["page_count"] = len(text_parts)
            result.source = f"{result.source}+pdf"

        return result


class DocxExtractor(BaseExtractor):
    name = "docx_extractor"

    def extract(self, raw_bytes: bytes, result: ScrapeResult, content_type: str = "") -> ScrapeResult:
        """Extract text from DOCX bytes, appending markdown to the result."""
        if not raw_bytes or content_type not in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
            ""
        ):
            return result

        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            paragraphs: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    if para.style and para.style.name and para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading ", "")
                        prefix = "#" * max(1, min(6, int(level) if level.isdigit() else 2))
                        paragraphs.append(f"{prefix} {text}")
                    else:
                        paragraphs.append(text)

            markdown = "\n\n".join(paragraphs)
            result.markdown = markdown
            result.text = "\n\n".join(doc.paragraphs)
            result.metadata["paragraph_count"] = len(doc.paragraphs)
            result.source = f"{result.source}+docx"
        except ImportError:
            pass

        return result
