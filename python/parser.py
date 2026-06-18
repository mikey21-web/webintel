"""Document parser — converts PDF/DOCX to clean markdown via URL or file upload."""

import io
import re
from typing import Optional, Tuple
from urllib.parse import urlparse

import httpx


def detect_document_type(url: str, content_type: Optional[str] = None) -> Optional[str]:
    """Detect document format from URL extension or content-type header."""
    if content_type:
        ct = content_type.lower().split(";")[0].strip()
        if ct == "application/pdf":
            return "pdf"
        if ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                  "application/msword"):
            return "docx"

    ext = url.split("?")[0].split("#")[0].rsplit(".", 1)[-1].lower()
    if ext in ("pdf",):
        return "pdf"
    if ext in ("docx", "doc"):
        return "docx"

    return None


async def fetch_document(url: str) -> Tuple[bytes, str, str]:
    """Fetch a document from a URL. Returns (raw_bytes, content_type, detected_format)."""
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
        "Accept": "application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document,application/msword,*/*",
    }

    async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
        resp = await client.get(url, headers=headers)
        resp.raise_for_status()
        ct = resp.headers.get("content-type", "")
        doc_type = detect_document_type(url, ct)
        if not doc_type:
            raise ValueError(f"Unsupported document type: {ct}. URL: {url}")
        return resp.content, ct, doc_type


def pdf_to_markdown(raw_bytes: bytes) -> Tuple[str, dict]:
    """Convert PDF bytes to markdown. Returns (markdown, metadata)."""
    text_parts: list[str] = []
    metadata: dict = {"parser": "none"}

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            metadata["page_count"] = len(pdf.pages)
            metadata["parser"] = "pdfplumber"
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except ImportError:
        pass

    if not text_parts:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
            metadata["page_count"] = len(reader.pages)
            metadata["parser"] = "pypdf2"
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        except ImportError:
            pass

    if text_parts:
        return "\n\n".join(text_parts), metadata

    return "", {"error": "No PDF parser available (install pdfplumber or PyPDF2)", "parser": "none"}


def docx_to_markdown(raw_bytes: bytes) -> Tuple[str, dict]:
    """Convert DOCX bytes to markdown. Returns (markdown, metadata)."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(raw_bytes))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                depth = max(1, min(6, int(level) if level.isdigit() else 2))
                paragraphs.append(f"{'#' * depth} {text}")
            else:
                paragraphs.append(text)

        return "\n\n".join(paragraphs), {
            "parser": "python-docx",
            "paragraph_count": len(doc.paragraphs),
        }
    except ImportError:
        return "", {"error": "python-docx not installed", "parser": "none"}


def is_document_url(url: str) -> bool:
    """Quick check if a URL likely points to a document."""
    return detect_document_type(url) is not None
